"""生成 法务 review 简报 Word 文档, 给律师用.

输出: D:\\Batch\\chroniccare\\docs\\LEGAL_REVIEW_BRIEF.docx
"""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- helpers ----------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    return p


def add_code_block(doc, text):
    """添加等宽字体块,作为 md 原文引用."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        # 中文字体 fallback
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_number(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ''
            cell.paragraphs[0].add_run(str(val)).font.size = Pt(10)
    return t


def add_hr(doc):
    p = doc.add_paragraph()
    run = p.add_run('─' * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


# ---------- 文档 ----------

doc = Document()

# 全局默认字体
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(10.5)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ===== 封面 =====
add_heading(doc, '慢病管家(ChronicCare)法务 Review 简报', 0)
add_para(doc, '—— 3 份法律文档 + 12 个 P0 问题 + 律师交付物', italic=True, size=11)
doc.add_paragraph()

meta_table = doc.add_table(rows=5, cols=2)
meta_table.style = 'Light Shading Accent 1'
meta = [
    ('文档日期', '2026-08-02'),
    ('项目版本', 'v0.27.0+64 (round 82)'),
    ('产品', '慢病管家 / ChronicCare —— Flutter 端精神心理 / 慢病管理 App'),
    ('待上架', 'Apple App Store + Google Play (中国大陆 / 港澳台 / 英文 3 locale)'),
    ('联系方式', 'support@chroniccare.app (待注册)'),
]
for i, (k, v) in enumerate(meta):
    meta_table.rows[i].cells[0].text = k
    meta_table.rows[i].cells[1].text = v
    for run in meta_table.rows[i].cells[0].paragraphs[0].runs:
        run.bold = True
doc.add_paragraph()

# ===== 一、项目简介 =====
add_heading(doc, '一、项目简介(给律师 5 分钟可读)', 1)

add_para(doc, '产品:慢病管家(ChronicCare)—— Flutter 端精神心理 / 慢病管理 App')
add_para(doc, '核心功能:')
add_bullet(doc, '每日服药打卡 + 提醒')
add_bullet(doc, '心理评估(PHQ-9 / GAD-7)')
add_bullet(doc, '每日情绪记录')
add_bullet(doc, '私密倾诉空间(树洞,文字 + 录音)')
add_bullet(doc, '失联通知(连续多日未打卡 → 自动 SMS 通知预设紧急联系人)')
add_bullet(doc, '数据本地导出 / 导入(JSON)')

add_heading(doc, '1.1 技术架构核心(影响合规判断)', 2)
add_para(doc, '零云端:所有数据存用户设备本地,不上传任何服务器。')
add_para(doc, '数据库:SQLCipher AES-256 加密(密钥由设备 SecureRandom 生成,绑定设备)。')
add_para(doc, '录音:AES-256 加密(密钥与数据库密钥独立管理,存 SecureStorage)。')
add_para(doc, '树洞文字:AES-256 字段级加密(v0.21 起)。')

add_heading(doc, '1.2 当前 release 模式(v0.27.0+64) 3 个业务暂停', 2)
add_para(doc, '工程当前有 3 项业务暂停,需法务评估文案 / 实际一致性:', bold=True)
add_number(doc, '失联通知整体暂停:`FeatureFlags.emergencyContactEnabled = false`,`AliyunSmsProvider.send()` 仍 throw UnimplementedError。用户可加联系人但实际不触发任何通知。')
add_number(doc, 'IAP 8 元买断暂停:`FeatureFlags._prodIapEnabled = false`,App 内无购买入口。但用户协议 §3 仍写"售价 8 元,一次性买断"。')
add_number(doc, '隐私 / PIPL 投诉邮箱软隐藏:`privacy@chroniccare.app` 邮箱不实际接收,用户通过 App 内"设置 → 法律与隐私"页行使撤回同意权。')

add_heading(doc, '1.3 法务 review 重点', 2)
add_para(doc, '3 份法律 md 都是 R67-R69 工程团队自写,未经律师过审(每份 md 末尾"修订历史"段都标"草稿 (未经律师过审)")。')
add_para(doc, '计划 2026-08-15 提交 Apple App Store + Google Play(中国大陆 + 港澳台 + 英文 3 个 locale)。')
add_para(doc, '目标 1.0 上线,不可压缩瓶颈 = 律师签字 + 法务过审。')

doc.add_page_break()

# ===== 二、3 份法律文档 =====
add_heading(doc, '二、3 份法律文档(请律师逐字 review)', 1)
add_para(doc, '以下 3 份 md 是工程团队自写版本,每份末尾"修订历史"段都标"草稿"。请律师逐字 review 并出修订意见。', italic=True)

# 2.1 隐私政策
add_heading(doc, '2.1 隐私政策(privacy_policy.md,14.5 KB,224 行)', 2)

privacy_policy = '''# 隐私政策

> 依据:《中华人民共和国个人信息保护法》(PIPL)、《App 违法违规收集使用个人信息行为认定方法》

## 0. 同意记录

依据《PIPL》§14 个人信息处理者处理敏感个人信息前应取得"单独同意"。

App 在首次启动时要求用户分别勾选 3 项同意:

1. 《用户协议》(通用条款)
2. 《隐私政策》(本文)
3. 《敏感个人信息处理同意书》(健康医疗等敏感数据)

3 项同意的时刻 + 协议版本号会写入本地数据库(可审计,不可篡改):

- `userAgreementVersion`(如 `v0.22-2026-07-21`)
- `privacyPolicyVersion`
- `sensitiveDataConsentAt`(同意时间)
- `consentRevokedAt`(撤回时间,null = 未撤回)

## 0.5 紧急联系人告知

依据《PIPL》§23 个人信息处理者向第三方(紧急联系人)提供个人信息前,应取得个人**单独同意**且**告知第三方**。

App 在首次设置时要求用户勾选"我已告知上述联系人,App 会在我失联时给他们发通知"才允许进入下一步。该勾选 + 联系人列表会作为失联通知触发的合法依据。

> **v0.27 round 66 (R66) 软提示更新**: R66 起在 setup 流程对每个联系人弹**软提示**("我已告知 TA 我会发送失联通知"勾选 + 单独告知弹窗), 不再是单一"全局勾选"。"已告知"软提示作为失联通知触发的合法依据 (PIPL §23 间接证据)。联系人本人未回复 Y 确认 (R55 TODO — 依赖 SMS provider 真接) 期间, 失联通知**整体业务暂停** (`FeatureFlags.emergencyContactEnabled = false`, R66 双层防御), 触发链路不发送任何数据给第三方。

## 1. 我们收集哪些信息

| 信息类别 | 具体字段 | 收集目的 | 存储位置 | 敏感 |
|---|---|---|---|---|
| 用户标识 | 用户昵称(可填可不填) | 在通知 / 报告中"我是 XXX" | 本地加密数据库 (SQLCipher) | 否 |
| 紧急联系人 | 姓名 + 手机号 | 失联通知用 | 本地加密数据库 | 是(第三方 PII) |
| 健康数据 | 药名、剂量、打卡时间、PHQ-9 / GAD-7 评分 | 服药追踪、趋势分析 | 本地加密数据库 | **是** (PIPL §28 健康医疗) |
| 私密倾诉 | 文字 / 录音(树洞) | 用户主动记录 | **本地加密**:文字 + 录音均使用 AES-256 加密(密钥设备绑定) | **是** (PIPL §28 私密记录) |
| 设备信息 | 设备型号、操作系统版本 | 仅本地判断通知兼容性,不存储不上传 | 不收集 | 否 |

**我们不收集:**位置、通讯录、相册、相机(录音除外)、设备 ID、广告 ID。

## 2. 信息的存储与保护

- 所有数据**仅存在用户设备本地**,**不上传任何云端**
- 数据库使用 **SQLCipher AES-256** 加密(密钥由设备 SecureRandom 生成,绑定设备)
- 数据库密钥**不离开设备**,不存储在我们或任何第三方服务器
- 录音文件存放在 App 私有目录 (`/data/data/<package>/files/`),**v0.18 起录音使用 AES-256 加密**(密钥与数据库密钥独立管理)
- 树洞文字字段使用 **AES-256 加密存储**(v0.21 起),即便设备 root + 拿到 DB 加密 key,无 SecureStorage key 仍无法读字段明文

## 3. 信息的共享

**我们不共享任何用户数据给第三方**,除非:
- 用户**主动**使用"导出 JSON 备份"功能(数据保存在用户剪贴板或用户指定位置)
- 用户**主动**使用"分享"功能(数据通过用户选择的 App 分享)
- 法律法规要求(如公安 / 检察机关依法调取)

**失联通知触发时**,我们会将下列信息发送给用户**预设**的紧急联系人:
- 用户昵称
- 距上次打卡的天数
- 一条预设的关怀短信模板

**我们不主动向紧急联系人以外的任何人发送信息。**

> **v0.27 round 66 (R66) 现状**: 失联通知业务整体暂停 (`FeatureFlags.emergencyContactEnabled = false`)。本节描述的**全部**数据流 (用户昵称 / 距上次打卡天数 / 关怀短信模板 → 紧急联系人) 在本版本**不实际触发**。设置 → 紧急联系人入口仅作为用户**预先配置**, 等 v0.28 真接 SMS provider + 完成 PIPL §38 跨境评估后启用。

## 4. 用户的权利

依据 PIPL,您享有以下权利:

| 权利 | 行使方式 |
|---|---|
| 知情权 | 阅读本政策 |
| 决定权 | 选择是否提供某些信息(药名可填"匿名") |
| 查询权 | 在 App 内查看所有打卡 / 评估 / 树洞记录 |
| 更正权 | 在 App 内修改 / 删除任何记录 |
| 删除权 | 在 App 内删除单条 / 全部数据;卸载 App 立即清除所有本地数据 |
| 撤回同意 | 在「设置 → 法律与隐私」页随时撤回"失联通知 / 树洞 / 趋势分析"3 项可选功能;撤回后该功能立即停用,**数据不删除**(可重新开启),撤回时刻记录到本地供审计 |
| 注销权 | 卸载 App = 完全注销,所有本地数据立即清除 |

## 5. 文字与录音加密

- **录音文件 (v0.18 起)**:AES-256 加密,密钥存 FlutterSecureStorage(Android encryptedSharedPreferences / iOS Keychain),IV 每次随机
- **树洞文字 (v0.21 起)**:AES-256 字段级加密,同样存 SecureStorage 密钥
- **加密范围**:设备 root + 拿到 DB 加密 key + 拿到 SecureStorage key 三者凑齐仍能解密(这是加密的本质——保护"没拿到 key 的人")
- **未加密部分**:打卡记录、联系人、用药、评估、情绪等非敏感字段(数据库整体 SQLCipher 加密已足够)

## 6. Cookie 与追踪

本 App **不使用任何 Cookie、追踪 SDK、广告 SDK**。

## 7. 第三方依赖

本 App 使用的第三方库(均不收集用户数据):

- `flutter_secure_storage` — 密钥存储(iOS Keychain / Android Keystore)
- `sqlcipher_flutter_libs` — 数据库加密
- `flutter_local_notifications` — 本地通知
- `audioplayers` / `record` — 录音 / 播放
- `go_router` / `riverpod` / `drift` — 框架

完整列表见 `pubspec.yaml`。

## 8. 政策的变更

- 重大变更会通过 App 内通知告知
- 继续使用本 App 视为接受变更

## 9. 联系方式

- 个人信息保护负责人:**本服务不提供邮件渠道** (v0.27 R67 Sprint 1 决策, 软隐藏 `privacy@chroniccare.app`)
  用户可通过 **App 内 设置 → 法律与隐私** 页面行使 PIPL §14 撤回同意权 (R67 ConsentGate 集中器统一执行, 撤回后业务立即停止)
  详见 `docs/LEGACY_API_NOTES.md` 了解软隐藏决策 + 重新启用条件
- 投诉举报:可向网信办、公安机关举报

## 10. 未成年人保护

依据《未成年人保护法》§44 与《个人信息保护法》§31,网络产品和服务对未成年人有特殊保护义务。

**本 App 的未成年人保护措施:**

- 本 App 设计面向**成年人(18 周岁以上)**使用,**不建议 14 周岁以下儿童单独使用**(精神心理类健康数据为敏感 PII,需监护人共同阅读本政策并辅助使用)
- 14-18 周岁用户**需监护人代为签署同意**,App 在设置流程中要求勾选"我已满 18 周岁 或 已由监护人代为同意"
- App 不会主动收集任何 14 周岁以下用户的敏感个人信息
- 如发现 14 周岁以下用户误用,监护人可通过 **App 内 设置 → 法律与隐私** 页面发起数据删除请求,我们将在 7 个工作日内响应 (v0.27 R67 Sprint 1 决策, 软隐藏 `privacy@chroniccare.app` 邮件渠道)
- 本 App 不会向未成年人推送任何营销内容、广告或诱导付费信息

## 11. 跨境数据传输 (v0.27 R54 增补, R69 版本号 walkthrough)

依据《个人信息保护法》§38、§39、§40 与《数据出境安全评估办法》,
个人信息处理者向境外提供个人信息需满足以下任一条件:

1. 通过国家网信部门组织的安全评估
2. 经专业机构进行个人信息保护认证
3. 按照国家网信部门制定的标准合同与境外接收方订立协议
4. 法律、行政法规或者国家网信部门规定的其他条件

**本 App 跨境数据传输场景(失联通知):**

当用户配置的紧急联系人**手机号归属地为境外**(非 +86 大陆号段)
时,失联通知 SMS / Email 触发的数据传输链路涉及**跨境 PII 传输**:

| 字段 | 接收方 | 跨境链路 |
|------|--------|----------|
| 用户昵称 | 紧急联系人本人 | App → SMS provider → 境外运营商 → 紧急联系人 |
| 距上次打卡天数 | 紧急联系人本人 | 同上 |
| 关怀短信模板 | 紧急联系人本人 | 同上 |

**本 App 的跨境 PII 保护措施:**

- **失联通知触发前,App 必须在设置流程中要求用户对每个境外紧急联系人单独勾选"我已告知 TA 我会发送失联通知"**(PIPL §23 单独告知第三方)。
- **跨境传输前,App 显示告知弹窗**,明确说明:接收方所在国家/地区、传输的 PII 字段、传输目的(PII 紧急联系)、接收方仅用于关怀。
- **跨境 SMS provider 选型时优先选择境内有 PIPL 跨境合规备案的 provider**(如阿里云国际版、Twilio + 境内主体)。
- **用户可随时撤回境外紧急联系人** —— App 设置页提供"移除联系人"操作,撤回后该联系人不再接收失联通知,跨境传输立即停止。
- **跨境 PII 传输审计日志**(本地)记录:传输时间、接收方手机号、传输字段、SMS provider 名称、用户 ID。供 PIPL §54 审计。

**未涉及跨境场景(无需 PIPL §38 跨境评估):**

- App 数据库全部存放在用户设备本地,无云端服务器,无数据出境
- 失联通知仅发给用户在 App 内**主动添加**的紧急联系人
- 树洞 / 健康数据 / 联系人列表**永不上传**到任何服务器

**法务声明:**

> v0.27 R69 walkthrough: 失联通知业务整体暂停 (`FeatureFlags.emergencyContactEnabled=false`),
> release 模式下 `AliyunSmsProvider.send()` 仍 throw UnimplementedError (R55+ 占位)。
> R68 commit `d691551` 修了 CareEngine safety consent 撤回业务层真接,
> use case `FireCareStrategyUseCase` 入口 `if (isSafetyConsentWithdrawn) → disabled` 早返。
>
> **正式上 store 前必须:**(1) 接境内备案的 SMS provider; (2) 完成 PIPL §38
> 跨境评估或标准合同备案; (3) 法务 review 本政策。

## 12. 紧急联系人"单独同意"实现进度 (PIPL §13)

依据《个人信息保护法》§13、§14、§23、§29,处理敏感个人信息前应
取得**单独同意**,向第三方提供 PII 前应**单独告知第三方**。

**当前实现状态 (v0.27 R69):**

| 阶段 | 实现细节 | 完成度 |
|------|----------|--------|
| 用户本人同意 | 设置流程要求勾选 3 项法律协议 (用户协议 / 隐私政策 / 敏感数据同意书) | ✅ v0.22 |
| 紧急联系人"已告知"勾选 | 设置流程要求勾选"我已告知上述联系人" | ✅ v0.22 |
| 紧急联系人**软告知弹窗** | R66 起对每个联系人单独弹"我已告知 TA 我会发送失联通知"软提示 | ✅ v0.27 R66 |
| 紧急联系人回复 Y 确认 | 短信回复确认机制 | ⏸ v0.27 业务整体暂停 (FeatureFlags.emergencyContactEnabled=false, 依赖 SMS provider 真接 — R55+ 占位, R28+ 启用) |
| 同意记录可审计 | 写入 `user_profiles.consent_*` 字段 | ✅ v0.21 |
| 撤回同意 | 设置页"法律与隐私"页可撤回 3 项功能 | ✅ v0.22 |
| 撤回同意**业务层生效** | vent_repository / CareEngine / trend_page 真的拦截 | ✅ v0.27 R67 (Sprint 1) |

**法律风险(完成前):** 严格按 PIPL §13/§23,联系人本人未确认 = 单独告知
未真正完成,失联通知 SMS 触发的合法性依赖用户勾选作为"间接证据"。
**风险等级:** 中(国内 4 store 上架审核可能打回,法务风险存在但量刑轻)。

**修复路径:** v0.28 (规划中) 接 SMS provider 后,在 setup 流程对每个联系人
发短信"我是 XXX,我已设置你为我的紧急联系人,如我失联会发通知,回复 Y 确认"。
联系人回复 Y 后写 `user_profiles.contact_consent_confirmed_at` 字段。
联系人 30 天未回复 Y 视为未同意,失联通知不发(graceful degrade)。'''

add_code_block(doc, privacy_policy)

doc.add_page_break()

# 2.2 用户协议
add_heading(doc, '2.2 用户协议(user_agreement.md,4.6 KB,83 行)', 2)

user_agreement = '''# 用户协议

## 1. 服务说明

「慢病管家」(以下简称"本 App")是一款面向慢性病、精神心理疾病患者及其家属的健康管理工具,核心功能包括:

- 每日服药打卡提醒
- 用药记录与历史趋势
- 心理评估(PHQ-9 / GAD-7)
- 私密倾诉空间(树洞)
- **失联通知**(**规划中,本版本未启用** — 连续多日未打卡时,自动通知预设的紧急联系人)
- 邮件 / 短信关怀通知
- 数据本地备份与恢复(导出/导入 JSON)

## 2. 服务范围

- **本 App 不提供医疗建议、诊断或治疗**。所有用药提醒、心理评估量表仅供参考,不能替代专业医师面诊。
- 失联通知功能**不是紧急救援服务**。遇紧急情况请拨打 120 / 110 或联系最近医院。

## 3. 付费规则

本 App 售价人民币 8 元(Google Play / Apple App Store 统一定价),一次性买断,**不收取订阅费**。

> **v0.27 R69 更新**: 当前 release 模式 IAP 业务整体暂停
> (`FeatureFlags._prodIapEnabled = false`, R68 commit `d691551` 决策),
> v0.27 本版本 App 内**不显示**"立即买断"入口。本付费规则段在
> **v0.28 真接 productId 后启用**,届时用户协议版本号 bump 后
> 重新走用户同意流程刷 `userAgreementVersion` 字段。

## 4. 退款政策

- Google Play / Apple App Store 按其平台规则处理退款申请
- 购买后 24 小时内未使用核心功能(每日打卡)可全额退款
- 已使用核心功能后,如遇重大 Bug 影响使用,可联系开发者协商

## 5. 免责声明

- 本 App 不对以下情况负责:
  - 因用户未设置紧急联系人导致失联时无人通知
  - 因 SMS 通道未连接(默认 mock 状态)导致通知未发出
  - 因用户误填药名 / 剂量导致医疗差错
  - 因设备丢失 / 损坏导致本地数据丢失(请定期导出 JSON 备份)
  - 因 iOS / Android 系统升级导致 App 暂时不可用
- 心理评估结果仅供参考,不应作为临床诊断依据。如有严重心理困扰,请联系当地精神卫生中心或心理危机干预热线(**北京 010-82951332 / 全国 400-161-9995**)。

## 6. 用户行为规范

- 不得利用本 App 骚扰他人(包括但不限于:把失联通知当成恶意骚扰工具)
- 不得利用本 App 进行任何违法违规活动
- 不得对本 App 进行反向工程、破解或分发盗版

## 7. 协议变更

- 本 App 保留随时修改本协议的权利
- 重大变更会通过 App 内通知 + 应用商店更新说明告知
- 继续使用本 App 视为接受变更后的协议

## 8. 联系方式

- 开发者邮箱:`support@chroniccare.app`(**TODO 占位 — 上 store 前必须注册并替换为真实邮箱**, 详见 `docs/SPRINT1_LEGAL_TODO.md`)
- GitHub Issues:`https://github.com/example/chroniccare/issues`(**TODO 占位,需确认或替换为真实项目仓库**, 详见 `docs/SPRINT1_LEGAL_TODO.md`)
- 隐私 / PIPL 投诉:**本服务不提供邮件渠道** (v0.27 R67 Sprint 1 决策, 软隐藏 `privacy@chroniccare.app`)
  用户可通过 **App 内 设置 → 法律与隐私** 页面行使 PIPL §14 撤回同意权 (R67 ConsentGate 集中器统一执行, 撤回后业务立即停止)
  详见 `docs/LEGACY_API_NOTES.md` 了解软隐藏决策 + 重新启用条件'''

add_code_block(doc, user_agreement)

doc.add_page_break()

# 2.3 敏感信息同意书
add_heading(doc, '2.3 敏感个人信息处理同意书(sensitive_data_consent.md,4.7 KB,106 行)', 2)

sensitive_consent = '''# 敏感个人信息处理同意书

> 依据:《中华人民共和国个人信息保护法》(PIPL) 第 28-29 条关于"敏感个人信息"的规定

## 1. 什么是敏感个人信息

依据 PIPL,**敏感个人信息**是指一旦泄露或者非法使用,容易导致自然人的人格尊严受到侵害或者人身、财产安全受到危害的个人信息,包括:

- 医疗健康信息
- 生物识别信息
- 宗教信仰
- 特定身份
- 金融账户
- 行踪轨迹
- 14 周岁以下未成年人的信息

## 2. 本 App 处理的敏感个人信息

本 App 处理以下敏感个人信息:

### 2.1 健康医疗信息

- **药名、剂量、用药时间** — 用于服药追踪
- **打卡时间戳** — 用于失联检测(**规划中,本版本未启用**)
- **PHQ-9 / GAD-7 心理评估答案** — 用于评估抑郁 / 焦虑症状
- **每日心情评分** — 用于情绪趋势

### 2.2 私密倾诉内容(树洞)

- **树洞文字** — 用户主动记录
- **树洞录音** — 用户主动录制

> ⚠️ 树洞是**最高敏感**的数据。我们对此的承诺:
> - 树洞内容**绝不**进入趋势 / 评估 / 失联检测 / 通知 / 任何统计
> - 树洞内容**绝不**离开用户设备
> - 树洞数据**不**包含在"导出 JSON"中(只导出文字,录音因路径问题无法跨设备复用)
> - 树洞数据导入后,会作为新条目添加(不会覆盖现有树洞)

## 3. 处理目的与方式

| 敏感信息 | 处理目的 | 处理方式 | 存储位置 |
|---|---|---|---|
| 药名 + 剂量 | 服药追踪 / 趋势 | 本地写入 + 本地查询 | 本地加密 DB |
| 打卡时间 | 失联检测(**规划中**) | 本地计算 | 本地加密 DB |
| 心理评估答案 | 自评症状严重程度 | 本地评分 + 本地查询 | 本地加密 DB |
| 树洞文字 | 用户主动记录 | 本地写入 | 本地加密 DB |
| 树洞录音 | 用户主动记录 | 本地文件 | **本地加密存储(AES-256,密钥设备绑定,2026-07 起启用)** |

## 4. 您的明确同意

依据 PIPL,处理敏感个人信息**必须取得您的单独同意**,而不能"勾选总协议即视为全部同意"。

因此本 App 单独设置此同意书。您需要在 setup 流程中**主动勾选**(默认不勾选)才能使用本 App。

您可以**随时撤回**同意,撤回方式:

- 在设置页关闭"失联通知"功能(**规划中,本版本未启用**)(撤回"紧急联系人手机号"的使用同意)
- 卸载 App(撤回全部同意,所有数据立即清除)
- 在 App 内删除单条 / 全部敏感数据

## 5. 不提供敏感信息的后果

- **不提供姓名**:失联通知(**规划中**)无法个性化("XXX 已 3 天未打卡" → "用户已 3 天未打卡")
- **不提供紧急联系人**:失联通知功能无法启用
- **不提供药名**:用药追踪只能记录时间,无法记录具体用药
- **不做心理评估**:无法看到评估历史趋势
- **不使用树洞**:不影响其他功能

## 6. 处理期限

- 用户主动删除的数据:立即清除
- 用户卸载 App:所有本地数据立即清除
- 长期保留:用户可永久保留数据,只要 App 还在使用

**我们不会因为任何理由在服务器保留您的数据(因为我们没有服务器)。**

## 7. 您的特别权利(敏感信息)

| 权利 | 行使方式 |
|---|---|
| 单独同意 | 在 setup 流程单独勾选本同意书 |
| 单独撤回 | 在设置页单独关闭"失联通知"等功能(**规划中,本版本未启用**) |
| 单独删除 | 在 App 内删除单条记录 |
| 单独导出 | 在设置页"导出 JSON"单独导出 |'''

add_code_block(doc, sensitive_consent)

doc.add_page_break()

# ===== 三、12 个 P0 问题 =====
add_heading(doc, '三、请律师回答的 12 个 P0 问题', 1)
add_para(doc, '12 个 P0 问题按风险等级排序,前 4 个最高风险(Q1-Q3 涉及合规红线、Q4-Q5 涉及数据导出 / 跨境)。', italic=True)

questions = [
    {
        'num': 'Q1',
        'title': '失联通知业务暂停的合规性',
        'context': '隐私政策 §0.5/§3/§12 描述"失联通知会发 SMS 给紧急联系人",但 FeatureFlags.emergencyContactEnabled = false 实际不触发任何通知。',
        'subs': [
            'Q1a:"规划中,本版本未启用"的措辞是否够强?能否在 Apple 5.2.1 / Google Play 4.3 审核 + 监管询问时站得住?',
            'Q1b:是否需要在主页 / 紧急联系人 section 顶部加显眼 banner?',
            'Q1c:v0.27 release 不开 / v0.28 接入 SMS 后开的切换,文案是否需统一升级?',
        ],
    },
    {
        'num': 'Q2',
        'title': '紧急联系人"单独同意"链路不完整',
        'context': '严格按 PIPL §23"向第三方提供 PII 前应单独告知第三方",未回复 Y = 未真正单独告知 = 中度合规风险。隐私政策自评"风险等级:中"。',
        'subs': [
            'Q2a:在业务暂停期间,R66 的"软告知"是否能作为合规抗辩?',
            'Q2b:是否需要完全移除"失联通知"在所有渠道的描述,只写"未来规划"?',
        ],
    },
    {
        'num': 'Q3',
        'title': '8 元买断 IAP 暂停 + 文档写明',
        'context': '用户协议 §3 写"售价 8 元,一次性买断",但 release 模式无购买入口。',
        'subs': [
            'Q3a:二选一 — 删 §3 8 元买断段(推到 v0.28)或启用 IAP(推荐)?法务建议哪个?',
            'Q3b:若删,需不需要保留"8 元"承诺在历史 user_agreement 中(让老用户能查阅)?',
        ],
    },
    {
        'num': 'Q4',
        'title': '数据导出 0 consent 流程',
        'context': 'R82 工程已实现导出走 ConsentDialog + audit log,但法务需确认。',
        'subs': [
            'Q4a:隐私政策需不需要补一段说明"导出 = 数据可携权行使,PIPL §44 留痕"?',
            'Q4b:导出 JSON 是明文,法务是否要求强制加密 + 提示风险?',
            'Q4c:用户导出 → 误传到云盘 → 被第三方抓取的责任边界?(需不需要"导出后用户自担风险"声明?)',
        ],
    },
    {
        'num': 'Q5',
        'title': '跨境数据传输 PIPL §38',
        'context': '业务暂停期间,隐私政策 §11 已描述跨境场景。',
        'subs': [
            'Q5a:业务暂停期间,隐私政策 §11 描述跨境场景是否合适(还是等 v0.28 真接时再加)?',
            'Q5b:v0.28 接 SMS 后,境内用户加境外联系人的前置告知弹窗强度是否够?',
            'Q5c:§38 评估是"上架前"做还是"接 SMS 前"做?(建议接 SMS 前)',
        ],
    },
    {
        'num': 'Q6',
        'title': '隐私 / PIPL 投诉邮箱软隐藏',
        'context': 'PIPL §50 应提供便捷的权利行使渠道 + 明确的联系方式。软隐藏邮箱是否合规?',
        'subs': [
            'Q6a:必须注册 privacy@chroniccare.app 邮箱真接收 + 7 工作日响应?',
            'Q6b:App 内"法律与隐私"页是否够"便捷"?(需测用户操作流程)',
            'Q6c:未成年人 7 工作日响应承诺 vs 软隐藏邮箱,实际能响应吗?',
        ],
    },
    {
        'num': 'Q7',
        'title': '树洞录音是否"敏感个人信息"',
        'context': '树洞(精神心理类表达)在 PIPL §28 定义中归类待确认。',
        'subs': [
            'Q7a:树洞在 PIPL §28 敏感个人信息定义中属于"医疗健康"还是"私密记录"?',
            'Q7b:撤回 ConsentKind.vent 同意后,旧树洞不删除(可重新开启)是否合规 PIPL §47?',
            'Q7c:AES-256 加密 + SecureStorage 密钥管理是否满足 PIPL §51 加密要求?',
        ],
    },
    {
        'num': 'Q8',
        'title': '第三方 SDK 披露不完整',
        'context': '隐私政策 §7 列了 6 个 SDK,但 pubspec.yaml 实际 16 个依赖。缺 in_app_purchase / speech_to_text / pdf / printing / permission_handler 等。',
        'subs': [
            'Q8a:第三方 SDK 列表谁负责核?(法务 vs 工程)',
            'Q8b:SDK 收集数据是否真"零"?(比如 in_app_purchase 必收购买历史)如何呈现?',
            'Q8c:是否需要每个 SDK 单独签订数据处理协议(DPA)?',
        ],
    },
    {
        'num': 'Q9',
        'title': '自动化决策 / 算法透明性(PIPL §24)',
        'context': '失联检测算法走 care_engine.dart(用户连续 N 天未打卡 + 各种 strategy 评估),这是典型"自动化决策"。',
        'subs': [
            'Q9a:隐私政策是否需加"算法逻辑"说明?',
            'Q9b:用户拒绝方式(撤回同意)是否够便捷?',
        ],
    },
    {
        'num': 'Q10',
        'title': '免责声明边界',
        'context': '精神心理类 App 需在每个相关页面 / 评估结果页都提醒"非诊断依据"。',
        'subs': [
            'Q10a:用户协议 §5 措辞是否覆盖全部精神心理相关功能(PHQ-9 / GAD-7 / 情绪趋势 / 失联通知)?',
            'Q10b:4 国危机热线是否完整?(中国 / 美国 / 英国 / 国际 — 缺台湾 1925 / 香港 2389 2222 / 澳门?)',
        ],
    },
    {
        'num': 'Q11',
        'title': '未成年人 14-18 周岁验证',
        'context': 'App 设置流程要求勾选"我已满 18 周岁 或 已由监护人代为同意"。',
        'subs': [
            'Q11a:"勾选"是否够?还是需要身份证 OCR / 活体检测?',
            'Q11b:14-18 周岁"监护人代为同意"如何留存证据?(PIPL §14 单独同意需"明确")',
            'Q11c:7 工作日响应机制 vs 软隐藏邮箱,实际能响应吗?',
        ],
    },
    {
        'num': 'Q12',
        'title': 'IAP 自动续费 + 平台抽成',
        'context': 'Apple 强制 IAP 抽成 30%(8 元 → 开发者 5.6 元),Google Play 30% / 15%。',
        'subs': [
            'Q12a:8 元买断是否合规中国"明码标价"规定?',
            'Q12b:退款政策(用户协议 §4:24 小时内未使用核心功能可全额退款)是否符合 Apple / Google 平台规则 + 中国消法?',
            'Q12c:是否需考虑国内应用市场(IAP 抽成不同)的差异化定价?',
        ],
    },
]

for q in questions:
    add_heading(doc, f'{q["num"]}. {q["title"]}', 2)
    add_quote(doc, f'【现状】{q["context"]}')
    add_para(doc, '请律师回答以下子问题:', bold=True)
    for sub in q['subs']:
        p = doc.add_paragraph(sub, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

doc.add_page_break()

# ===== 四、律师交付物 =====
add_heading(doc, '四、律师交付物(请律师给 4 样)', 1)
add_para(doc, '期望律师在 1-2 周内交付以下 4 样材料:', bold=True)

add_number(doc, '修订跟踪表(Word/Excel)— 每份 md 每个段落,标 ✅ 通过 / ⚠️ 修改(给建议)/ ❌ 删除')
add_number(doc, '签字盖章版 PDF(3 份:隐私政策 / 用户协议 / 敏感个人信息处理同意书)')
add_number(doc, '法务声明信 — 3 份 md 都已经过律师 review,符合 PIPL / 中国相关法规')
add_number(doc, '风险提示清单 — 上面 12 个 P0 问题,每项律师标注 ✅ 可接受 / ⚠️ 需修订 / ❌ 必改')

doc.add_page_break()

# ===== 五、估时 + 费用 =====
add_heading(doc, '五、估时 + 费用', 1)
add_table(doc, ['项目', '数值'], [
    ('律师费用', '¥15-30k / 文档 × 3 文档 = ¥45-90k'),
    ('Review 时间', '1-2 周'),
    ('关键瓶颈', '不修 1.0 没法正式发;不可压缩'),
    ('律师专业方向', 'PIPL / 精神心理类 App / App 隐私政策 / 医疗免责声明'),
    ('建议律所', '锦天城 / 中伦 / 大成 律所网络法务团队'),
])

add_heading(doc, '5.1 上 store 前必完成清单(给律师对齐时间)', 2)
add_para(doc, '法务 review 1-2 周期间,以下事项可并行启动:', bold=True)
add_number(doc, '注册 chroniccare.app 域名 + ICP 备案(7-20 天,跟律师并行)')
add_number(doc, '生成 Release keystore(1-2h,工程操作)')
add_number(doc, 'macOS 跑 pod install + 截 33 个 App Store 截图 + App Icon(1-2 天,需 macOS + 设计师)')
add_number(doc, 'App Store Connect + Google Play Console 配置(半天,需 Apple ID / 阿里云 / 域名)')

add_para(doc, '5 项中只有"法务 review"是不可压缩瓶颈,其他都可并行。')

doc.add_page_break()

# ===== 附录:项目技术信息 =====
add_heading(doc, '附录:项目技术信息(给律师参考)', 1)

add_heading(doc, '附.1 关键代码文件', 2)
add_para(doc, '律师 review 时如需查证技术实现细节,可参考以下文件:')
add_bullet(doc, 'lib/core/data/feature_flags.dart — 当前 3 个业务暂停的开关')
add_bullet(doc, 'lib/domain/entities/consent_artifact.dart — 同意基础实体(5 个 ConsentKind)')
add_bullet(doc, 'lib/core/l10n/strings.dart — domain 层 i18n 集中器')
add_bullet(doc, 'lib/presentation/pages/setup/setup_page.dart — setup 流程(3 文件同意)')
add_bullet(doc, 'lib/presentation/pages/contact/contacts_list_widget.dart — 联系人 + ConsentDialog 软提示')
add_bullet(doc, 'lib/presentation/pages/settings/legal_page.dart — 设置 → 法律与隐私(撤回同意)')

add_heading(doc, '附.2 集中器文档', 2)
add_bullet(doc, 'docs/SPRINT1_LEGAL_TODO.md — Sprint 1 法务 todo 集中器')
add_bullet(doc, 'docs/LEGACY_API_NOTES.md — 软隐藏邮箱决策')
add_bullet(doc, 'docs/VERSION_1.0_PLAN.md — 1.0 上架时间线 + 法务门槛')
add_bullet(doc, 'docs/STOREFRONT_RELEASE_SOP.md — R82 上架前手动 checklist')

add_hr(doc)

add_para(doc, '本文档由工程团队(Mavis / MiniMax Code)自动生成,基于 v0.27 round 82 状态。', italic=True, size=9)
add_para(doc, '生成时间:2026-08-02 | 工程 commit:3fcf4f4 (legal brief) / a634dc9 (上架冲刺 A)', italic=True, size=9)
add_para(doc, '后续 v0.28 / v1.0 重大版本更新(失联通知真接 SMS / IAP 真接 / 5 厂商 push 等)需重新过法务。', italic=True, size=9)

# 保存
output = io.open('docs/LEGAL_REVIEW_BRIEF.docx', 'wb')
doc.save(output)
output.close()
print('OK: docs/LEGAL_REVIEW_BRIEF.docx')
